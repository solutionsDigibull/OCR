"""DOCX file processor."""

import os
from typing import Dict, Any

from .base import BaseProcessor
from ..result import ConversionResult
from ..exceptions import ConversionError, FileNotFoundError


class DOCXProcessor(BaseProcessor):
    """Processor for Microsoft Word DOCX and DOC files."""
    
    def can_process(self, file_path: str) -> bool:
        """Check if this processor can handle the given file.
        
        Args:
            file_path: Path to the file to check
            
        Returns:
            True if this processor can handle the file
        """
        if not os.path.exists(file_path):
            return False
        
        # Check file extension - ensure file_path is a string
        file_path_str = str(file_path)
        _, ext = os.path.splitext(file_path_str.lower())
        return ext in ['.docx', '.doc']
    
    def process(self, file_path: str) -> ConversionResult:
        """Process the DOCX file and return a conversion result.
        
        Args:
            file_path: Path to the DOCX file to process
            
        Returns:
            ConversionResult containing the processed content
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            ConversionError: If processing fails
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Initialize metadata
        metadata = {
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "processor": "DOCXProcessor"
        }
        
        # Check file extension - ensure file_path is a string
        file_path_str = str(file_path)
        _, ext = os.path.splitext(file_path_str.lower())
        
        if ext == '.doc':
            return self._process_doc_file(file_path, metadata)
        else:
            return self._process_docx_file(file_path, metadata)
    
    def _process_doc_file(self, file_path: str, metadata: Dict[str, Any]) -> ConversionResult:
        """Process .doc files using pypandoc."""
        try:
            import pypandoc
            
            # Convert .doc to markdown using pandoc
            content = pypandoc.convert_file(file_path, 'markdown')
            
            metadata.update({
                "file_type": "doc",
                "extractor": "pypandoc"
            })
            
            # Clean up the content
            content = self._clean_content(content)
            
            return ConversionResult(content, metadata)
            
        except ImportError:
            raise ConversionError("pypandoc is required for .doc file processing. Install it with: pip install pypandoc")
        except Exception as e:
            raise ConversionError(f"Failed to process .doc file {file_path}: {str(e)}")
    
    def _process_docx_file(self, file_path: str, metadata: Dict[str, Any]) -> ConversionResult:
        """Process .docx files using python-docx with improved table extraction."""
        try:
            from docx import Document

            content_parts = []
            doc = Document(file_path)

            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "section_count": len(doc.sections),
                "file_type": "docx",
                "extractor": "python-docx"
            })

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if this is a heading
                    if paragraph.style.name.startswith('Heading'):
                        level = paragraph.style.name.replace('Heading ', '')
                        try:
                            level_num = int(level)
                            content_parts.append(f"\n{'#' * min(level_num, 6)} {paragraph.text}\n")
                        except ValueError:
                            content_parts.append(f"\n## {paragraph.text}\n")
                    else:
                        content_parts.append(paragraph.text)

            # Extract text from tables (improved)
            for table_idx, table in enumerate(doc.tables):
                # Check if preserve_layout is available (from base class or config)
                preserve_layout = getattr(self, 'preserve_layout', False)
                if preserve_layout:
                    content_parts.append(f"\n### Table {table_idx+1}\n")

                # Gather all rows
                rows = table.rows
                if not rows:
                    continue

                # Detect merged cells (optional warning)
                merged_warning = False
                for row in rows:
                    for cell in row.cells:
                        if len(cell._tc.xpath('.//w:vMerge')) > 0 or len(cell._tc.xpath('.//w:gridSpan')) > 0:
                            merged_warning = True
                            break
                    if merged_warning:
                        break
                if merged_warning:
                    content_parts.append("*Warning: Table contains merged cells which may not render correctly in markdown.*\n")

                # Row limit for large tables
                row_limit = 20
                if len(rows) > row_limit:
                    content_parts.append(f"*Table truncated to first {row_limit} rows out of {len(rows)} total.*\n")

                # Build table data
                table_data = []
                for i, row in enumerate(rows):
                    if i >= row_limit:
                        break
                    row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    table_data.append(row_data)

                # Ensure all rows have the same number of columns
                max_cols = max(len(r) for r in table_data)
                for r in table_data:
                    while len(r) < max_cols:
                        r.append("")

                # Markdown table: first row as header
                if table_data:
                    header = table_data[0]
                    separator = ["---"] * len(header)
                    content_parts.append("| " + " | ".join(header) + " |")
                    content_parts.append("| " + " | ".join(separator) + " |")
                    for row in table_data[1:]:
                        content_parts.append("| " + " | ".join(row) + " |")
                    content_parts.append("")

            content = '\n'.join(content_parts)
            content = self._clean_content(content)
            return ConversionResult(content, metadata)
        except ImportError:
            raise ConversionError("python-docx is required for .docx file processing. Install it with: pip install python-docx")
        except Exception as e:
            raise ConversionError(f"Failed to process .docx file {file_path}: {str(e)}")
    
    def _clean_content(self, content: str) -> str:
        """Clean up the extracted Word content.
        
        Args:
            content: Raw Word text content
            
        Returns:
            Cleaned text content
        """
        # Remove excessive whitespace and normalize
        lines = content.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove excessive whitespace
            line = ' '.join(line.split())
            if line.strip():
                cleaned_lines.append(line)
        
        # Join lines and add proper spacing
        content = '\n'.join(cleaned_lines)
        
        # Add spacing around headers
        content = content.replace('## ', '\n## ')
        content = content.replace('### ', '\n### ')
        
        return content.strip() 